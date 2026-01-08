"""
–ü—Ä–µ–¥—Å—Ç–∞–≤–ª–µ–Ω–∏–µ –¥–ª—è –ø–µ—á–∞—Ç–∏ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ (—É–ø—Ä–æ—â—ë–Ω–Ω—ã–π –∏–Ω—Ç–µ—Ä—Ñ–µ–π—Å –¥–ª—è –ø–æ–∂–∏–ª—ã—Ö)
"""

import os
from typing import Optional
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QFrame, QMessageBox, QTextEdit
)
from PyQt6.QtCore import Qt, pyqtSignal, QThread, pyqtSlot
from PyQt6.QtGui import QPixmap, QImage, QFont
from PIL import Image
import fitz  # PyMuPDF for PDF preview

from .styles import Styles
from .file_picker_dialog import FilePickerDialog
from .print_settings_dialog import PrintSettingsDialog
from .print_confirmation_dialog import PrintConfirmationDialog
from ..models import PrintSettings
from ..services import PrinterService, ImageProcessingService, logger
from ..services.sound_service import sound_service
from ..services.settings_storage import settings_storage

# –ü–æ–ø—ã—Ç–∫–∞ –∏–º–ø–æ—Ä—Ç–∞ python-docx –¥–ª—è –ø–æ–¥–¥–µ—Ä–∂–∫–∏ Word –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORTED = True
except ImportError:
    DOCX_SUPPORTED = False
    logger.warning("python-docx –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω, –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä DOCX –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω")


class PrintWorker(QThread):
    """–†–∞–±–æ—á–∏–π –ø–æ—Ç–æ–∫ –¥–ª—è –ø–µ—á–∞—Ç–∏"""
    finished = pyqtSignal(bool, str)

    def __init__(self, printer_service: PrinterService, file_path: str, settings: PrintSettings):
        super().__init__()
        self.printer_service = printer_service
        self.file_path = file_path
        self.settings = settings

    def run(self):
        try:
            self.printer_service.print_file(self.file_path, self.settings)
            self.finished.emit(True, "–ü–µ—á–∞—Ç—å —É—Å–ø–µ—à–Ω–æ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω–∞")
        except Exception as e:
            logger.exception(f"–û—à–∏–±–∫–∞ –ø–µ—á–∞—Ç–∏: {e}")
            self.finished.emit(False, str(e))


class PrintView(QWidget):
    """–ü—Ä–µ–¥—Å—Ç–∞–≤–ª–µ–Ω–∏–µ –¥–ª—è –ø–µ—á–∞—Ç–∏ —Å —É–ø—Ä–æ—â—ë–Ω–Ω—ã–º –∏–Ω—Ç–µ—Ä—Ñ–µ–π—Å–æ–º"""

    navigate_back = pyqtSignal()

    def __init__(self, printer_service: PrinterService, image_processing: ImageProcessingService, parent=None):
        super().__init__(parent)
        self._printer_service = printer_service
        self._image_processing = image_processing
        self._current_file: Optional[str] = None
        self._pdf_document = None
        self._current_page = 0
        self._total_pages = 1
        self._original_image: Optional[Image.Image] = None
        self._settings = PrintSettings()
        self._print_worker: Optional[PrintWorker] = None
        self._docx_text: Optional[str] = None

        self._init_ui()
        logger.info("–û—Ç–∫—Ä—ã—Ç–∞ —Å—Ç—Ä–∞–Ω–∏—Ü–∞ –ø–µ—á–∞—Ç–∏")

    def _init_ui(self):
        """–ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è –∏–Ω—Ç–µ—Ä—Ñ–µ–π—Å–∞"""
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(30, 25, 30, 25)
        main_layout.setSpacing(20)

        # –ó–∞–≥–æ–ª–æ–≤–æ–∫
        header_layout = QHBoxLayout()

        back_btn = QPushButton("‚Üê –ù–∞–∑–∞–¥")
        back_btn.setFixedSize(150, 60)
        back_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {Styles.TEXT_SECONDARY};
                color: white;
                font-size: {Styles.FONT_SIZE_NORMAL}px;
                font-weight: bold;
                border-radius: 10px;
            }}
            QPushButton:hover {{ background-color: #5a6268; }}
        """)
        back_btn.clicked.connect(self.navigate_back.emit)
        header_layout.addWidget(back_btn)

        header_layout.addStretch()

        title_label = QLabel("üñ®Ô∏è –ü–ï–ß–ê–¢–¨")
        title_font = QFont()
        title_font.setPointSize(Styles.FONT_SIZE_TITLE)
        title_font.setBold(True)
        title_label.setFont(title_font)
        header_layout.addWidget(title_label)

        header_layout.addStretch()
        header_layout.addSpacing(150)

        main_layout.addLayout(header_layout)

        # –û—Å–Ω–æ–≤–Ω–æ–π –∫–æ–Ω—Ç–µ–Ω—Ç
        content_layout = QHBoxLayout()
        content_layout.setSpacing(25)

        # –õ–µ–≤–∞—è –ø–∞–Ω–µ–ª—å - –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä
        preview_panel = self._create_preview_panel()
        content_layout.addWidget(preview_panel, stretch=2)

        # –ü—Ä–∞–≤–∞—è –ø–∞–Ω–µ–ª—å - –¥–µ–π—Å—Ç–≤–∏—è
        actions_panel = self._create_actions_panel()
        content_layout.addWidget(actions_panel, stretch=1)

        main_layout.addLayout(content_layout)

        # –ö–Ω–æ–ø–∫–∞ –ø–µ—á–∞—Ç–∏
        self._print_btn = QPushButton("üñ®Ô∏è  –ù–ê–ü–ï–ß–ê–¢–ê–¢–¨")
        self._print_btn.setFixedHeight(100)
        self._print_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {Styles.SUCCESS_COLOR};
                color: white;
                font-size: {Styles.FONT_SIZE_TITLE}px;
                font-weight: bold;
                border-radius: 15px;
            }}
            QPushButton:hover {{
                background-color: #43A047;
            }}
            QPushButton:disabled {{
                background-color: #BDBDBD;
            }}
        """)
        self._print_btn.setEnabled(False)
        self._print_btn.clicked.connect(self._on_print_clicked)
        main_layout.addWidget(self._print_btn)

    def _create_preview_panel(self) -> QFrame:
        """–°–æ–∑–¥–∞—Ç—å –ø–∞–Ω–µ–ª—å –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä–∞"""
        frame = QFrame()
        frame.setStyleSheet(Styles.get_card_style())
        layout = QVBoxLayout(frame)
        layout.setSpacing(15)

        # –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ —Ñ–∞–π–ª–µ
        self._file_info_label = QLabel("–§–∞–π–ª –Ω–µ –≤—ã–±—Ä–∞–Ω")
        self._file_info_label.setStyleSheet(f"""
            font-size: {Styles.FONT_SIZE_LARGE}px;
            color: {Styles.TEXT_SECONDARY};
            padding: 10px;
        """)
        self._file_info_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self._file_info_label)

        # –û–±–ª–∞—Å—Ç—å –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä–∞
        preview_container = QFrame()
        preview_container.setStyleSheet("background-color: #F5F5F5; border-radius: 12px;")
        preview_layout = QVBoxLayout(preview_container)

        # Label –¥–ª—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π/PDF
        self._preview_label = QLabel("–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–∞–π–ª –¥–ª—è –ø–µ—á–∞—Ç–∏")
        self._preview_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._preview_label.setStyleSheet(f"""
            color: {Styles.TEXT_SECONDARY};
            font-size: {Styles.FONT_SIZE_LARGE}px;
        """)
        self._preview_label.setMinimumSize(400, 450)
        self._preview_label.setScaledContents(False)
        preview_layout.addWidget(self._preview_label)

        # TextEdit –¥–ª—è DOCX –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä–∞
        self._docx_preview = QTextEdit()
        self._docx_preview.setReadOnly(True)
        self._docx_preview.setVisible(False)
        self._docx_preview.setStyleSheet(f"""
            QTextEdit {{
                background-color: white;
                border: none;
                padding: 20px;
                font-family: 'Times New Roman', serif;
                font-size: {Styles.FONT_SIZE_NORMAL}px;
            }}
        """)
        preview_layout.addWidget(self._docx_preview)

        layout.addWidget(preview_container, stretch=1)

        # –ù–∞–≤–∏–≥–∞—Ü–∏—è –ø–æ —Å—Ç—Ä–∞–Ω–∏—Ü–∞–º
        self._nav_widget = QWidget()
        nav_layout = QHBoxLayout(self._nav_widget)
        nav_layout.setContentsMargins(0, 10, 0, 0)

        self._prev_btn = QPushButton("‚óÄ –ù–∞–∑–∞–¥")
        self._prev_btn.setFixedSize(120, 50)
        self._prev_btn.setStyleSheet(f"font-size: {Styles.FONT_SIZE_NORMAL}px;")
        self._prev_btn.clicked.connect(self._prev_page)
        nav_layout.addStretch()
        nav_layout.addWidget(self._prev_btn)

        self._page_label = QLabel("–°—Ç—Ä–∞–Ω–∏—Ü–∞ 1 –∏–∑ 1")
        self._page_label.setStyleSheet(f"font-size: {Styles.FONT_SIZE_LARGE}px; margin: 0 20px;")
        nav_layout.addWidget(self._page_label)

        self._next_btn = QPushButton("–í–ø–µ—Ä—ë–¥ ‚ñ∂")
        self._next_btn.setFixedSize(120, 50)
        self._next_btn.setStyleSheet(f"font-size: {Styles.FONT_SIZE_NORMAL}px;")
        self._next_btn.clicked.connect(self._next_page)
        nav_layout.addWidget(self._next_btn)
        nav_layout.addStretch()

        self._nav_widget.setVisible(False)
        layout.addWidget(self._nav_widget)

        return frame

    def _create_actions_panel(self) -> QFrame:
        """–°–æ–∑–¥–∞—Ç—å –ø–∞–Ω–µ–ª—å –¥–µ–π—Å—Ç–≤–∏–π"""
        frame = QFrame()
        frame.setStyleSheet(Styles.get_card_style())
        layout = QVBoxLayout(frame)
        layout.setSpacing(20)
        layout.setContentsMargins(25, 25, 25, 25)

        # –ó–∞–≥–æ–ª–æ–≤–æ–∫
        title = QLabel("–í—ã–±–µ—Ä–∏—Ç–µ —Ñ–∞–π–ª")
        title.setStyleSheet(f"""
            font-size: {Styles.FONT_SIZE_LARGE}px;
            font-weight: bold;
            color: {Styles.TEXT_PRIMARY};
        """)
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)

        layout.addSpacing(10)

        # –ö–Ω–æ–ø–∫–∞ –≤—ã–±–æ—Ä–∞ —Ñ–∞–π–ª–∞
        select_btn = QPushButton("üìÇ  –í—ã–±—Ä–∞—Ç—å —Ñ–∞–π–ª")
        select_btn.setFixedHeight(80)
        select_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {Styles.PRIMARY_COLOR};
                color: white;
                font-size: {Styles.FONT_SIZE_LARGE}px;
                font-weight: bold;
                border-radius: 12px;
            }}
            QPushButton:hover {{
                background-color: #1976D2;
            }}
        """)
        select_btn.clicked.connect(self._on_select_file_clicked)
        layout.addWidget(select_btn)

        layout.addSpacing(20)

        # –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ –∫–æ–ø–∏—è—Ö
        self._copies_label = QLabel("–ö–æ–ø–∏–π: 1")
        self._copies_label.setStyleSheet(f"""
            font-size: {Styles.FONT_SIZE_LARGE}px;
            color: {Styles.TEXT_PRIMARY};
            padding: 15px;
            background-color: #F5F5F5;
            border-radius: 8px;
        """)
        self._copies_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self._copies_label)

        # –ö–Ω–æ–ø–∫–∞ –Ω–∞—Å—Ç—Ä–æ–µ–∫ –ø–µ—á–∞—Ç–∏
        settings_btn = QPushButton("‚öôÔ∏è  –ù–∞—Å—Ç—Ä–æ–π–∫–∏ –ø–µ—á–∞—Ç–∏")
        settings_btn.setFixedHeight(70)
        settings_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {Styles.TEXT_SECONDARY};
                color: white;
                font-size: {Styles.FONT_SIZE_NORMAL}px;
                font-weight: bold;
                border-radius: 10px;
            }}
            QPushButton:hover {{
                background-color: #5a6268;
            }}
        """)
        settings_btn.clicked.connect(self._on_settings_clicked)
        layout.addWidget(settings_btn)

        layout.addStretch()

        # –ü–æ–¥—Å–∫–∞–∑–∫–∞
        hint_label = QLabel("–ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ–º—ã–µ —Ñ–æ—Ä–º–∞—Ç—ã:\nPDF, Word, –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è")
        hint_label.setStyleSheet(f"""
            color: {Styles.TEXT_SECONDARY};
            font-size: {Styles.FONT_SIZE_NORMAL}px;
        """)
        hint_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(hint_label)

        return frame

    def _on_select_file_clicked(self):
        """–û—Ç–∫—Ä—ã—Ç—å –¥–∏–∞–ª–æ–≥ –≤—ã–±–æ—Ä–∞ —Ñ–∞–π–ª–∞"""
        dialog = FilePickerDialog(self)
        if dialog.exec():
            file_path = dialog.get_selected_file()
            if file_path:
                self._load_file(file_path)

    def _on_settings_clicked(self):
        """–û—Ç–∫—Ä—ã—Ç—å –¥–∏–∞–ª–æ–≥ –Ω–∞—Å—Ç—Ä–æ–µ–∫ –ø–µ—á–∞—Ç–∏"""
        dialog = PrintSettingsDialog(self._settings, self)
        if dialog.exec():
            self._settings = dialog.get_settings()
            self._copies_label.setText(f"–ö–æ–ø–∏–π: {self._settings.copies}")
            logger.info(f"–ù–∞—Å—Ç—Ä–æ–π–∫–∏ –ø–µ—á–∞—Ç–∏ –æ–±–Ω–æ–≤–ª–µ–Ω—ã: –∫–æ–ø–∏–π={self._settings.copies}")

    def load_file_for_print(self, file_path: str):
        """–ó–∞–≥—Ä—É–∑–∏—Ç—å —Ñ–∞–π–ª –¥–ª—è –ø–µ—á–∞—Ç–∏ (–≤—ã–∑—ã–≤–∞–µ—Ç—Å—è –∏–∑ drag-and-drop)"""
        self._load_file(file_path)

    def _load_file(self, file_path: str):
        """–ó–∞–≥—Ä—É–∑–∏—Ç—å —Ñ–∞–π–ª –¥–ª—è –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä–∞"""
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "–û—à–∏–±–∫–∞", f"–§–∞–π–ª –Ω–µ –Ω–∞–π–¥–µ–Ω:\n{file_path}")
            return

        self._current_file = file_path
        file_name = os.path.basename(file_path)
        self._file_info_label.setText(f"üìÑ {file_name}")
        logger.info(f"–ó–∞–≥—Ä—É–∂–µ–Ω —Ñ–∞–π–ª: {file_path}")

        # –î–æ–±–∞–≤–ª—è–µ–º –≤ –Ω–µ–¥–∞–≤–Ω–∏–µ —Ñ–∞–π–ª—ã
        settings_storage.add_recent_file(file_path)

        ext = os.path.splitext(file_path)[1].lower()

        # –°–∫—Ä—ã–≤–∞–µ–º –≤—Å–µ –≤–∏–¥–∂–µ—Ç—ã –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä–∞
        self._preview_label.setVisible(False)
        self._docx_preview.setVisible(False)

        if ext == '.pdf':
            self._load_pdf(file_path)
        elif ext in ('.docx', '.doc'):
            self._load_docx(file_path)
        else:
            self._load_image(file_path)

        self._print_btn.setEnabled(True)

    def _load_pdf(self, file_path: str):
        """–ó–∞–≥—Ä—É–∑–∏—Ç—å PDF –¥–ª—è –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä–∞"""
        try:
            if self._pdf_document:
                self._pdf_document.close()
            self._pdf_document = fitz.open(file_path)
            self._total_pages = len(self._pdf_document)
            self._current_page = 0
            self._docx_text = None
            self._original_image = None
            self._nav_widget.setVisible(self._total_pages > 1)
            self._preview_label.setVisible(True)
            self._update_page_info()
            self._render_pdf_page()
            logger.info(f"PDF –∑–∞–≥—Ä—É–∂–µ–Ω: {self._total_pages} —Å—Ç—Ä–∞–Ω–∏—Ü")
        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –æ—Ç–∫—Ä—ã—Ç–∏—è PDF: {e}")
            self._preview_label.setVisible(True)
            self._preview_label.setText(f"–ù–µ —É–¥–∞–ª–æ—Å—å –æ—Ç–∫—Ä—ã—Ç—å PDF:\n{e}")

    def _load_docx(self, file_path: str):
        """–ó–∞–≥—Ä—É–∑–∏—Ç—å DOCX –¥–ª—è –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä–∞"""
        if not DOCX_SUPPORTED:
            self._preview_label.setVisible(True)
            self._preview_label.setText("–î–ª—è –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä–∞ Word\n—É—Å—Ç–∞–Ω–æ–≤–∏—Ç–µ python-docx")
            logger.warning("python-docx –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω")
            return

        try:
            doc = DocxDocument(file_path)
            if self._pdf_document:
                self._pdf_document.close()
            self._pdf_document = None
            self._original_image = None
            self._nav_widget.setVisible(False)

            # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–µ–∫—Å—Ç
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))

            self._docx_text = "\n".join(full_text)
            self._docx_preview.setPlainText(self._docx_text)
            self._docx_preview.setVisible(True)

            lines = self._docx_text.count('\n') + 1
            self._total_pages = max(1, lines // 50)

            logger.info(f"DOCX –∑–∞–≥—Ä—É–∂–µ–Ω: ~{self._total_pages} —Å—Ç—Ä–∞–Ω–∏—Ü")
        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –æ—Ç–∫—Ä—ã—Ç–∏—è DOCX: {e}")
            self._preview_label.setVisible(True)
            self._preview_label.setText(f"–û—à–∏–±–∫–∞ –æ—Ç–∫—Ä—ã—Ç–∏—è –¥–æ–∫—É–º–µ–Ω—Ç–∞:\n{e}")

    def _load_image(self, file_path: str):
        """–ó–∞–≥—Ä—É–∑–∏—Ç—å –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –¥–ª—è –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä–∞"""
        try:
            self._original_image = Image.open(file_path)
            if self._pdf_document:
                self._pdf_document.close()
            self._pdf_document = None
            self._docx_text = None
            self._total_pages = 1
            self._current_page = 0
            self._nav_widget.setVisible(False)
            self._preview_label.setVisible(True)
            self._update_preview_image()
            logger.info(f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –∑–∞–≥—Ä—É–∂–µ–Ω–æ: {self._original_image.size}")
        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –æ—Ç–∫—Ä—ã—Ç–∏—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è: {e}")
            self._preview_label.setVisible(True)
            self._preview_label.setText(f"–ù–µ —É–¥–∞–ª–æ—Å—å –æ—Ç–∫—Ä—ã—Ç—å –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ:\n{e}")

    def _render_pdf_page(self):
        """–û—Ç—Ä–∏—Å–æ–≤–∞—Ç—å —Ç–µ–∫—É—â—É—é —Å—Ç—Ä–∞–Ω–∏—Ü—É PDF"""
        if not self._pdf_document:
            return

        page = self._pdf_document[self._current_page]
        mat = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=mat)

        img = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format.Format_RGB888)
        pixmap = QPixmap.fromImage(img)

        scaled = pixmap.scaled(
            self._preview_label.size(),
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation
        )
        self._preview_label.setPixmap(scaled)

    def _update_preview_image(self):
        """–û–±–Ω–æ–≤–∏—Ç—å –ø—Ä–µ–¥–ø—Ä–æ—Å–º–æ—Ç—Ä –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è"""
        if not self._original_image:
            return

        image = self._original_image
        if image.mode == 'RGBA':
            data = image.tobytes('raw', 'RGBA')
            qimg = QImage(data, image.width, image.height, QImage.Format.Format_RGBA8888)
        else:
            image = image.convert('RGB')
            data = image.tobytes('raw', 'RGB')
            qimg = QImage(data, image.width, image.height, QImage.Format.Format_RGB888)

        pixmap = QPixmap.fromImage(qimg)
        scaled = pixmap.scaled(
            self._preview_label.size(),
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation
        )
        self._preview_label.setPixmap(scaled)

    def _prev_page(self):
        """–ü—Ä–µ–¥—ã–¥—É—â–∞—è —Å—Ç—Ä–∞–Ω–∏—Ü–∞"""
        if self._current_page > 0:
            self._current_page -= 1
            self._update_page_info()
            self._render_pdf_page()

    def _next_page(self):
        """–°–ª–µ–¥—É—é—â–∞—è —Å—Ç—Ä–∞–Ω–∏—Ü–∞"""
        if self._current_page < self._total_pages - 1:
            self._current_page += 1
            self._update_page_info()
            self._render_pdf_page()

    def _update_page_info(self):
        """–û–±–Ω–æ–≤–∏—Ç—å –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ —Å—Ç—Ä–∞–Ω–∏—Ü–∞—Ö"""
        self._page_label.setText(f"–°—Ç—Ä–∞–Ω–∏—Ü–∞ {self._current_page + 1} –∏–∑ {self._total_pages}")
        self._prev_btn.setEnabled(self._current_page > 0)
        self._next_btn.setEnabled(self._current_page < self._total_pages - 1)

    def _on_print_clicked(self):
        """–û–±—Ä–∞–±–æ—Ç—á–∏–∫ –Ω–∞–∂–∞—Ç–∏—è –∫–Ω–æ–ø–∫–∏ –ø–µ—á–∞—Ç–∏"""
        if not self._current_file:
            return

        # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –¥–∏–∞–ª–æ–≥ –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∏—è
        file_name = os.path.basename(self._current_file)
        dialog = PrintConfirmationDialog(file_name, self._settings, self)

        if dialog.exec():
            logger.info(f"–ü–µ—á–∞—Ç—å –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–µ–Ω–∞: {self._current_file}")
            self._start_printing()
        else:
            logger.info("–ü–µ—á–∞—Ç—å –æ—Ç–º–µ–Ω–µ–Ω–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª–µ–º")

    def _start_printing(self):
        """–ù–∞—á–∞—Ç—å –ø–µ—á–∞—Ç—å"""
        self._print_btn.setEnabled(False)
        self._print_btn.setText("‚è≥ –ü–µ—á–∞—Ç–∞–µ–º...")

        self._print_worker = PrintWorker(self._printer_service, self._current_file, self._settings)
        self._print_worker.finished.connect(self._on_print_finished)
        self._print_worker.start()

    @pyqtSlot(bool, str)
    def _on_print_finished(self, success: bool, message: str):
        """–û–±—Ä–∞–±–æ—Ç—á–∏–∫ –∑–∞–≤–µ—Ä—à–µ–Ω–∏—è –ø–µ—á–∞—Ç–∏"""
        self._print_btn.setEnabled(True)
        self._print_btn.setText("üñ®Ô∏è  –ù–ê–ü–ï–ß–ê–¢–ê–¢–¨")

        if success:
            logger.info("–ü–µ—á–∞—Ç—å —É—Å–ø–µ—à–Ω–æ –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω–∞")
            sound_service.play_success()

            msg = QMessageBox(self)
            msg.setWindowTitle("–ì–æ—Ç–æ–≤–æ!")
            msg.setText("–î–æ–∫—É–º–µ–Ω—Ç –æ—Ç–ø—Ä–∞–≤–ª–µ–Ω –Ω–∞ –ø–µ—á–∞—Ç—å")
            msg.setInformativeText("–ó–∞–±–µ—Ä–∏—Ç–µ —Ä–∞—Å–ø–µ—á–∞—Ç–∫—É –∏–∑ –ª–æ—Ç–∫–∞ –ø—Ä–∏–Ω—Ç–µ—Ä–∞")
            msg.setIcon(QMessageBox.Icon.Information)
            msg.setStandardButtons(QMessageBox.StandardButton.Ok)
            msg.setStyleSheet(f"""
                QMessageBox {{
                    font-size: {Styles.FONT_SIZE_LARGE}px;
                }}
                QMessageBox QLabel {{
                    font-size: {Styles.FONT_SIZE_LARGE}px;
                    min-width: 400px;
                }}
                QPushButton {{
                    font-size: {Styles.FONT_SIZE_NORMAL}px;
                    min-width: 150px;
                    min-height: 50px;
                }}
            """)
            msg.exec()
        else:
            logger.error(f"–û—à–∏–±–∫–∞ –ø–µ—á–∞—Ç–∏: {message}")
            sound_service.play_error()

            msg = QMessageBox(self)
            msg.setWindowTitle("–û—à–∏–±–∫–∞")
            msg.setText("–ù–µ —É–¥–∞–ª–æ—Å—å –Ω–∞–ø–µ—á–∞—Ç–∞—Ç—å –¥–æ–∫—É–º–µ–Ω—Ç")
            msg.setInformativeText(f"–û—à–∏–±–∫–∞: {message}")
            msg.setIcon(QMessageBox.Icon.Warning)
            msg.setStandardButtons(QMessageBox.StandardButton.Ok)
            msg.setStyleSheet(f"""
                QMessageBox {{
                    font-size: {Styles.FONT_SIZE_LARGE}px;
                }}
                QMessageBox QLabel {{
                    font-size: {Styles.FONT_SIZE_LARGE}px;
                    min-width: 400px;
                }}
                QPushButton {{
                    font-size: {Styles.FONT_SIZE_NORMAL}px;
                    min-width: 150px;
                    min-height: 50px;
                }}
            """)
            msg.exec()

    def closeEvent(self, event):
        """–û—á–∏—Å—Ç–∫–∞ —Ä–µ—Å—É—Ä—Å–æ–≤"""
        if self._pdf_document:
            self._pdf_document.close()
        super().closeEvent(event)
